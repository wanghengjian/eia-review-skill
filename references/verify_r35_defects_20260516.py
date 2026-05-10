#!/usr/bin/env python3
"""
R35 缺陷逐一核实脚本
对某次审查的全部缺陷按关键词在 DOCX 中搜索验证

用法:
    python verify_r35_defects.py <docx_path> <review_id>
    
输出: 每条缺陷的关键词命中数 + 属实/存疑/不实判定
"""
import sys
import re
from pathlib import Path
from docx import Document

def extract_text(docx_path):
    doc = Document(docx_path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)

def verify_defects(text, defects):
    """defects: list of (rule_id, severity, chapter, desc, kw_pattern)"""
    print(f"总字符: {len(text):,}\n")
    header = f"{'No':>4}  {'规则':10}  {'严重度':4}  {'章节':16}  {'结果':8}  命中"
    print(header)
    print("-" * 80)
    for i, (rule_id, sev, chapter, desc, kw) in enumerate(defects, 1):
        hits = re.findall(kw, text, re.IGNORECASE)
        if hits:
            verdict = "✓ 属实"
            note = f"{len(hits)}处"
        else:
            short = kw.split("|")[0][:15]
            partial = re.findall(short[:8], text, re.IGNORECASE)
            if partial:
                verdict = "? 存疑"
                note = "模糊"
            else:
                verdict = "✗ 不实"
                note = "未找到"
        print(f"{i:4d}  {rule_id:10}  {sev:4}  {chapter:16}  {verdict:8}  {note}")

if __name__ == "__main__":
    docx = sys.argv[1] if len(sys.argv) > 1 else "report.docx"
    text = extract_text(docx)
    # 示例 defects 列表（从 DB 查出来后填入）
    print("用法: python verify_r35_defects.py <docx_path>")

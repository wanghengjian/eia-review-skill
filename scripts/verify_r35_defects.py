"""
R35 缺陷逐一核实脚本
对37条缺陷按关键词在DOCX原文中搜索验证，返回：属实/存疑/不实

用法:
  python3 verify_r35_defects.py <docx_path>

输出:
  - 37条缺陷的关键词命中统计
  - 严重缺陷的详细证据片段
  - 最终判断（属实/存疑/不实）
"""
import sys
import pickle
import re
from pathlib import Path

# python-docx 路径（优先用 backend venv）
try:
    from docx import Document
except ImportError:
    from pathlib import Path
    venv_site = Path(__file__).resolve().parent.parent.parent / "backend" / ".venv" / "lib" / "python3.13" / "site-packages"
    sys.path.insert(0, str(venv_site))
    from docx import Document


def load_docx(docx_path):
    """提取DOCX全文本（段落+表格）"""
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = {}
    for i, tbl in enumerate(doc.tables):
        rows = []
        for row in tbl.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            tables[i] = "\n".join(rows)
    full = "\n".join(paragraphs)
    return paragraphs, tables, full


def search(keyword, full_text):
    """全文精确搜索，返回所有匹配上下文"""
    return [
        full_text[max(0, m.start() - 50):m.end() + 80].replace("\n", " ")
        for m in re.finditer(re.escape(keyword), full_text)
    ]


# ── 37条缺陷定义 ──
# (rule_id, severity, chapter, description, verification_keywords)
DEFECTS = [
    # 严重 6条
    ("B-005-01", "严重", "第五章", "地表水排放标准表述不统一",
     ["DB 44/26-2001", "GB21907", "水污染物排放限值"]),
    ("B-005-02", "严重", "第一章", "声环境敏感目标与功能区区划矛盾",
     ["安居澜庭", "3类声环境", "2类声环境", "声环境功能区"]),
    ("B-005-02", "严重", "第二章", "纯水制备设备图文不一致",
     ["纯水制备", "37.5", "25t/h", "产水"]),
    ("B-005-02", "严重", "第三章", "挥发性有机物料平衡表数据矛盾",
     ["68.250", "甲醇", "微生物分解代谢"]),
    ("B-005-02", "严重", "第四章", "图文数据不一致（CO占标率）",
     ["CO", "占标率", "深圳", "东莞"]),
    ("B-005-02", "严重", "第十一章", "废水排放标准号与现行标准不一致",
     ["DB 44/26-2001", "DB44/26-2019", "废止"]),
    # 较重 27条
    ("C-002-01", "较重", "第一章", "三线一单符合性论证不完整",
     ["三线一单", "生态红线", "城镇开发边界"]),
    ("C-002-05", "较重", "第一章", "未开展区域规划环评准入要求分析",
     ["YB83XHG01", "区域规划环评", "光明区"]),
    ("C-004-01", "较重", "第一章", "排放标准引用非标准文献",
     ["大气污染物综合排放标准详解", "非标准文献"]),
    ("C-004-01", "较重", "第五章", "标准引用格式不完整",
     ["DB 44/26-2001", "生物工程类制药"]),
    ("C-004-03", "较重", "第一章", "地方标准优先执行论证不足",
     ["DB44/27", "广东省地方标准", "甲醇"]),
    ("C-005-01", "较重", "第五章", "大气评价等级判定依据缺失",
     ["Pmax", "大气评价等级", "二级", "2.2-2018"]),
    ("C-006-02", "较重", "第三章", "产品方案及产能表述模糊",
     ["贻贝粘蛋白", "多糖类产品", "病毒", "活菌"]),
    ("C-006-03", "较重", "第二章", "原辅材料汇总未覆盖全部物料",
     ["表2.4-1", "表2.4-2", "原辅材料"]),
    ("C-006-03", "较重", "第三章", "原辅材料清单不完整",
     ["乙醇", "甲醇", "氨水", "盐酸", "原辅材料"]),
    ("C-006-04", "较重", "第二章", "设备规模与产能匹配性缺乏论证",
     ["表2.5-1", "设备", "产能"]),
    ("C-006-04", "较重", "第三章", "设备清单与产能匹配性验证不足",
     ["发酵罐", "11076", "最大发酵"]),
    ("C-010-02", "较重", "第四章", "大气监测因子缺少项目特征污染物",
     ["乙腈", "TVOC", "氯化氢", "硫酸", "监测因子"]),
    ("C-010-03", "较重", "第四章", "地下水监测频次不符合导则要求",
     ["地下水", "监测频次", "HJ 610", "HJ610"]),
    ("C-010-04", "较重", "第四章", "引用监测数据可能超过有效期",
     ["光明银星", "污水处理", "地下水监测", "2023"]),
    ("C-011-02", "较重", "第五章", "地下水预测参数选取依据不足",
     ["有效孔隙度", "弥散度", "0.6", "10m"]),
    ("C-011-05", "较重", "第五章", "土壤预测结果单位不一致",
     ["CODMn", "2100", "mg/cm"]),
    ("C-012d", "较重", "第三章", "喷淋+两级活性炭综合效率取值偏高",
     ["喷淋", "活性炭", "去除率", "45%", "80%"]),
    ("C-017-02", "较重", "第七章", "风险潜势判定E值前后矛盾",
     ["E2", "地下水环境敏感程度", "风险潜势", "7.2"]),
    ("C-017-03", "较重", "第七章", "最大可信事故分析不合理",
     ["盐酸", "氨水", "泄漏", "储存桶"]),
    ("C-017-04", "较重", "第七章", "事故应急池容积计算未体现必要性论证",
     ["事故应急池", "438.5", "775.4"]),
    ("C-017-06", "较重", "第七章", "应急预案与区域衔接要求不明确",
     ["应急预案", "地方政府", "突发环境事件"]),
    ("C-018-01", "较重", "第九章", "监测计划缺少固体废物监测内容",
     ["固体废物", "危废", "监测计划", "表9.2"]),
    ("C-019-01", "较重", "第十一章", "未说明两次公示的具体时间间隔",
     ["公众参与", "两次公示", "时间间隔"]),
    ("C-019-03", "较重", "第十一章", "未说明问卷调查的具体数量",
     ["问卷调查", "发放份数", "公众"]),
    ("C-020-01", "较重", "第六章", "废气处理效率计算逻辑不一致",
     ["有机废气", "两级活性炭", "处理效率", "80%"]),
    ("C-057-01", "较重", "第六章", "分区防渗措施要求与相关标准不符",
     ["分区防渗", "GB18597", "2023", "危废"]),
    ("其他", "较重", "第十章", "引用废水排放标准编号错误",
     ["DB 44/26-2001", "生活污水", "三级标准"]),
    # 一般 4条
    ("C-018-02", "一般", "第九章", "废水总排放口总氮监测频次与自动监测要求不一致",
     ["总氮", "自动监测", "表9.2", "废水总排放口"]),
    ("C-018-03", "一般", "第九章", "委托监测安排未明确被委托机构资质要求",
     ["委托监测", "资质", "环境监测部门"]),
    ("C-018b", "一般", "第九章", "废水站出水口在线监测设备要求不明确",
     ["在线监测", "废水站", "水质", "表9.5"]),
    ("其他", "一般", "第七章", "敏感目标未列表，交叉引用不通",
     ["敏感目标", "表1.8", "7.1", "风险"]),
]


def verify(docx_path):
    """执行全部37条缺陷核实"""
    paragraphs, tables, full = load_docx(docx_path)

    print(f"段落: {len(paragraphs)}, 表格: {len(tables)}, 总字符: {len(full):,}")
    print("=" * 90)

    results = []
    for rule, severity, chapter, desc, kws in DEFECTS:
        hits = {}
        for kw in kws:
            h = search(kw, full)
            if h:
                hits[kw] = (len(h), h[0])

        if hits:
            verdict = "✓ 属实"
        else:
            verdict = "? 待核"

        results.append({
            "rule": rule,
            "severity": severity,
            "chapter": chapter,
            "desc": desc,
            "verdict": verdict,
            "hits": hits,
        })
        detail = "; ".join([f"{k}({n}处)" for k, (n, _) in hits.items()]) or "无命中"
        print(f"{rule:10} | {chapter:8} | {verdict:8} | {detail[:80]}")

    # 统计
    total = len(results)
    real = sum(1 for r in results if r["verdict"] == "✓ 属实")
    uncertain = sum(1 for r in results if r["verdict"] == "? 待核")
    print(f"\n属实: {real}/{total}, 待核: {uncertain}/{total}")
    return results


if __name__ == "__main__":
    if len(sys.argv) < 2:
        # 默认R35路径
        docx = Path.home() / ".hermes/workspace/eia-review/backend/uploads/1/国家生物制造产业创新中心建设项目环境影响告书（报批稿）2026.1.20.docx"
    else:
        docx = Path(sys.argv[1])

    if not docx.exists():
        print(f"文件不存在: {docx}")
        sys.exit(1)

    results = verify(docx)

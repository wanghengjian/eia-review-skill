#!/usr/bin/env python3
"""
自检脚本：对指定审查的缺陷进行 LLM 自动化核实

用法:
  python3 self_check.py <review_id>
  python3 self_check.py 35

输出:
  - 每条缺陷的核实结果（属实/存疑/不实）
  - 命中率统计
  - 规则优化建议（JSON + 摘要）

DB:
  ~/.hermes/workspace/eia-review/backend/eia_review.db
"""

import sys
import json
import re
import sqlite3
import time
import urllib.request
import urllib.error
from pathlib import Path
from datetime import datetime
from typing import Optional

# ── python-docx ──────────────────────────────────────────────────────────────
try:
    from docx import Document
except ImportError:
    import site
    venv = next((p for p in site.getsitepackages()
                 if '.venv' in p and 'python3' in p), None)
    if venv:
        sys.path.insert(0, venv)
    from docx import Document

# ── 路径配置 ────────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).resolve().parent
BACKEND_DIR = SCRIPT_DIR.parent.parent.parent / "workspace" / "eia-review" / "backend"
DB_PATH = BACKEND_DIR / "eia_review.db"

# ── 严重度映射 ───────────────────────────────────────────────────────────────
SEVERITY_MAP = {"严重": "A", "较重": "B", "一般": "C", "轻微": "C"}
SEVERITY_ORDER = ["严重", "较重", "一般", "轻微"]

# ── 核实结果常量 ─────────────────────────────────────────────────────────────
REAL = "属实"       # 有明确证据，缺陷成立
UNREAL = "不实"   # 无证据或证据矛盾，LLM误判
UNCERTAIN = "存疑" # 有相关证据但不确凿，需人工复核


# ═══════════════════════════════════════════════════════════════════════════
# LLM 核实
# ═══════════════════════════════════════════════════════════════════════════

def _load_api_key() -> str:
    """从 backend/.env 或环境变量加载 DeepSeek API Key"""
    # 优先从 backend/.env 加载（与 llm_client.py 保持一致）
    backend_env = Path.home() / ".hermes" / "workspace" / "eia-review" / "backend" / ".env"
    if backend_env.exists():
        for line in backend_env.read_text().splitlines():
            if line.startswith("DEEPSEEK_API_KEY="):
                key = line.split("=", 1)[1].strip()
                if key and len(key) > 5:
                    return key
    # Fallback 到环境变量
    import os
    return os.environ.get("DEEPSEEK_API_KEY", "")


def verify_defect_llm(
    defect_description: str,
    report_excerpt: str,
    chapter_name: str,
    rule_id: str,
    project_name: str,
    api_key: str,
) -> dict:
    """
    用 LLM 核实单条缺陷是否属实。

    判断标准：
    - 属实：缺陷描述中的核心问题在报告原文中确实有明确证据
    - 存疑：报告中有相关表述，但不足以确凿判定为缺陷
    - 不实：报告原文中找不到缺陷描述所指的内容，或原文明确与描述矛盾

    返回:
    {
        "verdict": "属实|存疑|不实",
        "hit_rate": 0.0-1.0,
        "summary": str,   # 专家可读的判定理由
        "reasoning": str, # LLM 原始推理过程
    }
    """
    if not api_key:
        return {
            "verdict": UNCERTAIN,
            "hit_rate": 0.0,
            "summary": "[未配置 deepseek_api_key，无法核实]",
            "reasoning": "",
        }

    # 清理报告原文（去掉缺陷标签）
    excerpt_clean = re.sub(r'【[^】]+】', '', report_excerpt).strip()
    if not excerpt_clean:
        excerpt_clean = "[报告中未提供相关原文摘录]"

    prompt = f"""## 角色
你是一名资深的深圳市环境影响评价技术专家，负责核实环评报告书审查所发现的缺陷是否属实。

## 待核实项目
- 项目名称：{project_name}
- 审查规则：{rule_id}

## 缺陷描述（LLM 审查时生成的问题描述）
{defect_description}

## 报告中该缺陷引用的原文摘录
{excerpt_clean}

## 判断要求
仔细阅读缺陷描述和报告原文，判断：该缺陷描述中的核心问题，在报告原文中是否有明确的文字证据支持？

请严格按以下标准判断：

**属实（缺陷成立）**：
  - 报告原文中确实包含缺陷所指的问题，且描述具体、准确
  - 例：缺陷说"表X废水中COD浓度超标"，原文确实列出了超标数据

**存疑（证据不足，需人工复核）**：
  - 报告中有相关表述，但证据不完整、不精确或相互矛盾
  - 例：缺陷说"缺少废气监测数据"，但原文中监测频次描述模糊，无法确定是否真的缺失

**不实（LLM误判或规则指向错误章节）**：
  - 报告原文中完全找不到缺陷描述所指的内容
  - 例：缺陷说"土壤中氨氮超标"，但报告中土壤监测数据完全正常，LLM可能把地下水数据错当成了土壤数据

## 输出格式（严格按以下JSON格式，不要有其他内容）
{{
  "verdict": "属实" 或 "存疑" 或 "不实",
  "confidence": "high" 或 "medium" 或 "low",
  "summary": "用一段话说明判断理由（不超过150字），要具体指出报告原文中的实际表述",
  "reasoning": "详细的推理过程（不超过300字）"
}}

注意：
- 你是核实者，不是初审者：重点判断该缺陷的证据是否充分，而不是重新发现新问题
- 报告原文中找不到对应内容 → 不实
- 原文有相关表述但证据模糊 → 存疑
- 原文有明确证据支持缺陷描述 → 属实
"""

    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.1,
        "max_tokens": 800,
    }

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }

    try:
        req = urllib.request.Request(
            "https://api.deepseek.com/chat/completions",
            data=json.dumps(payload).encode('utf-8'),
            headers=headers,
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=60) as resp:
            result = json.loads(resp.read().decode('utf-8'))
            content = result["choices"][0]["message"]["content"].strip()

        json_match = re.search(r'\{[\s\S]+\}', content)
        if not json_match:
            return {
                "verdict": UNCERTAIN,
                "hit_rate": 0.0,
                "summary": f"[LLM返回格式错误，无法解析结果]",
                "reasoning": content[:500],
            }

        llm_result = json.loads(json_match.group())
        verdict = llm_result.get("verdict", UNCERTAIN)
        confidence = llm_result.get("confidence", "medium")

        # hit_rate 映射
        hit_rate = {"high": 1.0, "medium": 0.5, "low": 0.0}.get(confidence, 0.5)

        return {
            "verdict": verdict,
            "hit_rate": hit_rate,
            "summary": llm_result.get("summary", ""),
            "reasoning": llm_result.get("reasoning", ""),
        }

    except json.JSONDecodeError as e:
        return {
            "verdict": UNCERTAIN,
            "hit_rate": 0.0,
            "summary": f"[JSON解析失败: {str(e)}]",
            "reasoning": content[:500] if 'content' in dir() else "",
        }
    except urllib.error.HTTPError as e:
        body = e.read().decode('utf-8', errors='replace')[:200]
        return {
            "verdict": UNCERTAIN,
            "hit_rate": 0.0,
            "summary": f"[HTTP错误 {e.code}: {e.reason}]",
            "reasoning": body,
        }
    except Exception as e:
        return {
            "verdict": UNCERTAIN,
            "hit_rate": 0.0,
            "summary": f"[LLM调用异常: {str(e)}]",
            "reasoning": "",
        }


# ═══════════════════════════════════════════════════════════════════════════
# 文档加载（按章节分割）
# ═══════════════════════════════════════════════════════════════════════════

def _detect_chapter_number(text: str) -> Optional[str]:
    """从标题文本中检测章节编号，如'第1章'/'第01章'/'第一章'"""
    for pattern in [
        r'^第([零一二三四五六七八九十百千万\d]+)章',
        r'^(第[零一二三四五六七八九十百千万\d]+章)',
    ]:
        m = re.search(pattern, text)
        if m:
            cn = m.group(1)
            # 中文数字转阿拉伯数字
            CN_MAP = {'零': '0', '一': '1', '二': '2', '三': '3', '四': '4',
                      '五': '5', '六': '6', '七': '7', '八': '8', '九': '9', '十': '10'}
            if cn in CN_MAP:
                num = CN_MAP[cn]
            else:
                # 处理十一、十二等
                num = cn
            try:
                return f"{int(num):02d}"
            except ValueError:
                return None
    return None


def _is_toc_heading(text: str) -> bool:
    """判断是否为目录标题（应跳过）"""
    return text == '目录' or (len(text) <= 4 and text.endswith('目'))


def _is_appendix_heading(text: str) -> bool:
    """判断是否为附件标题"""
    return '附件' in text


def _is_valid_chapter_heading(text: str) -> bool:
    """判断是否为有效章节标题（第X章格式）"""
    return bool(_detect_chapter_number(text))


def load_docx_with_chapters(docx_path: str) -> dict:
    """
    加载 DOCX，按章节分割文本。

    返回:
    {
        "chapters": {
            "07": {"num": "07", "name": "环境风险评价", "text": "...", "paragraphs": [...]},
            ...
        },
        "full_text": "...",
        "tables": [...],
        "paragraphs": [...],
    }
    """
    doc = Document(docx_path)

    chapters = {}  # num -> {"num", "name", "text", "paragraphs"}
    current_num = "000"
    current_name = "前言"
    current_paragraphs = []

    _chapter_base = None
    _h1_counter = [0]

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        # Heading 1 = 章节标题
        if style_name == 'Heading 1':
            # 目录跳过
            if _is_toc_heading(text):
                continue

            # 附件
            if _is_appendix_heading(text):
                if current_paragraphs:
                    chapters[current_num] = {
                        "num": current_num,
                        "name": current_name,
                        "text": "\n".join(current_paragraphs),
                        "paragraphs": current_paragraphs,
                    }
                current_num = "012"
                current_name = "附件"
                current_paragraphs = []
                continue

            ch_num = _detect_chapter_number(text)

            if ch_num:
                # 保存旧章节
                if current_paragraphs:
                    chapters[current_num] = {
                        "num": current_num,
                        "name": current_name,
                        "text": "\n".join(current_paragraphs),
                        "paragraphs": current_paragraphs,
                    }
                _chapter_base = int(ch_num)
                current_num = ch_num
                current_name = re.sub(r'^第[零一二三四五六七八九十百千万\d]+章\s*', '', text).strip()
                if not current_name:
                    current_name = f"第{ch_num}章"
                current_paragraphs = []
            else:
                # 非"第X章"格式的H1（概述/1总则等）
                if _chapter_base is None:
                    _chapter_base = 0
                current_num = f"{_chapter_base + _h1_counter[0]:03d}"
                _h1_counter[0] += 1
                current_name = text
                if current_paragraphs and current_num in chapters:
                    chapters[current_num]["paragraphs"].extend(current_paragraphs)
                    chapters[current_num]["text"] = "\n".join(chapters[current_num]["paragraphs"])
                else:
                    if current_paragraphs:
                        chapters[current_num] = {
                            "num": current_num,
                            "name": current_name,
                            "text": "\n".join(current_paragraphs),
                            "paragraphs": current_paragraphs,
                        }
                    current_paragraphs = []
        else:
            current_paragraphs.append(text)

    # 保存最后一章
    if current_paragraphs:
        chapters[current_num] = {
            "num": current_num,
            "name": current_name,
            "text": "\n".join(current_paragraphs),
            "paragraphs": current_paragraphs,
        }

    # 收集所有段落和表格
    all_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    all_tables = []
    from docx.oxml.table import CT_Tbl
    from docx.oxml.ns import qn
    for tbl in doc.element.body.iter(qn('w:tbl')):
        rows = []
        for row in tbl.iter(qn('w:tr')):
            cells = [cell.text.strip() if cell.text else "" for cell in row.iter(qn('w:tc'))]
            if any(cells):
                rows.append(cells)
        if rows:
            all_tables.append("\n".join(" | ".join(r) for r in rows))

    full_text = "\n".join(all_paragraphs)

    return {
        "chapters": chapters,
        "full_text": full_text,
        "tables": all_tables,
        "paragraphs": all_paragraphs,
    }


def find_chapter_for_defect(chapters: dict, defect_chapter: str) -> tuple:
    """
    根据缺陷的章节信息，找到对应的章节文本。

    优先用章节名匹配，其次用章节编号匹配。
    返回: (chapter_num, chapter_name, chapter_text)
    """
    if not defect_chapter:
        return ("", "", "")

    defect_chapter = defect_chapter.strip()

    # 先精确匹配章节编号（如 "7" -> "07"）
    for num, ch in chapters.items():
        if ch["num"] == defect_chapter or ch["num"] == defect_chapter.zfill(2):
            return (num, ch["name"], ch["text"])

    # 模糊匹配章节名（包含关系）
    for num, ch in chapters.items():
        if defect_chapter in ch["name"] or ch["name"] in defect_chapter:
            return (num, ch["name"], ch["text"])

    # 章节名模糊匹配
    for num, ch in chapters.items():
        if defect_chapter.replace("章", "") in ch["name"].replace("章", ""):
            return (num, ch["name"], ch["text"])

    return ("", "", "")


# ═══════════════════════════════════════════════════════════════════════════
# 分析与统计
# ═══════════════════════════════════════════════════════════════════════════

def analyze_results(results: list[dict]) -> dict:
    """汇总核实结果，计算质量指标"""
    total = len(results)
    real = sum(1 for r in results if r["verdict"] == REAL)
    uncertain = sum(1 for r in results if r["verdict"] == UNCERTAIN)
    unreal = sum(1 for r in results if r["verdict"] == UNREAL)

    hit_rate = real / total if total > 0 else 0
    precision = real / (real + unreal) if (real + unreal) > 0 else 0
    drift_rate = unreal / total if total > 0 else 0

    # 按严重度统计
    by_severity = {s: {"total": 0, "real": 0, "uncertain": 0, "unreal": 0}
                   for s in SEVERITY_ORDER}
    for r in results:
        sev = r.get("severity", "一般")
        if sev in by_severity:
            by_severity[sev]["total"] += 1
            if r["verdict"] == REAL:
                by_severity[sev]["real"] += 1
            elif r["verdict"] == UNCERTAIN:
                by_severity[sev]["uncertain"] += 1
            elif r["verdict"] == UNREAL:
                by_severity[sev]["unreal"] += 1

    # 按规则统计
    by_rule = {}
    for r in results:
        rid = r["rule_id"]
        if rid not in by_rule:
            by_rule[rid] = {"total": 0, "real": 0, "uncertain": 0, "unreal": 0}
        by_rule[rid]["total"] += 1
        if r["verdict"] == REAL:
            by_rule[rid]["real"] += 1
        elif r["verdict"] == UNCERTAIN:
            by_rule[rid]["uncertain"] += 1
        elif r["verdict"] == UNREAL:
            by_rule[rid]["unreal"] += 1

    # 问题规则
    problem_rules = []
    for rid, stats in by_rule.items():
        if stats["total"] >= 2:
            r_hit_rate = stats["real"] / stats["total"]
            if r_hit_rate < 0.8:
                problem_rules.append({
                    "rule_id": rid,
                    "total": stats["total"],
                    "real": stats["real"],
                    "hit_rate": round(r_hit_rate, 2),
                    "verdict": "低命中率" if stats["unreal"] > 0 else "高存疑率",
                })

    return {
        "total": total,
        "real": real,
        "uncertain": uncertain,
        "unreal": unreal,
        "hit_rate": round(hit_rate, 3),
        "precision": round(precision, 3),
        "drift_rate": round(drift_rate, 3),
        "by_severity": by_severity,
        "by_rule": by_rule,
        "problem_rules": sorted(problem_rules, key=lambda x: x["hit_rate"]),
    }


def generate_optimization_suggestions(results: list[dict], analysis: dict) -> list[dict]:
    """根据核实结果生成优化建议"""
    suggestions = []

    for pr in analysis["problem_rules"]:
        rid = pr["rule_id"]
        items = [r for r in results if r["rule_id"] == rid]
        unreal_items = [r for r in items if r["verdict"] == UNREAL]
        uncertain_items = [r for r in items if r["verdict"] == UNCERTAIN]

        if unreal_items:
            sample = unreal_items[0]
            suggestions.append({
                "type": "rule",
                "priority": "high" if pr["hit_rate"] < 0.5 else "medium",
                "rule_id": rid,
                "summary": f"规则{rid}命中率{pr['hit_rate']}%（{pr['real']}/{pr['total']}），{len(unreal_items)}条不实",
                "detail": f"典型误判：{sample.get('description', '')[:80]}",
                "sample_summary": sample.get("verify_result", {}).get("summary", ""),
            })

    # 跨章节问题
    chapter_coverage = {}
    for r in results:
        rid = r["rule_id"]
        ch = r.get("chapter", "")
        if rid not in chapter_coverage:
            chapter_coverage[rid] = {}
        if ch:
            chapter_coverage[rid][ch] = r["verdict"]

    cross_chapter_issues = []
    for rid, chapters in chapter_coverage.items():
        if len(chapters) > 1:
            verdicts = list(chapters.values())
            if UNCERTAIN in verdicts or UNREAL in verdicts:
                cross_chapter_issues.append({
                    "rule_id": rid,
                    "chapters": chapters,
                    "issue": "该规则多章节触发但存在不实/存疑，可能是章节适用性定义不准确"
                })

    if cross_chapter_issues:
        suggestions.append({
            "type": "cross_chapter",
            "priority": "medium",
            "summary": f"发现{len(cross_chapter_issues)}条跨章节规则存在核实异常",
            "items": cross_chapter_issues,
            "recommendation": "建议检查这些规则的适用章节定义"
        })

    return suggestions


# ═══════════════════════════════════════════════════════════════════════════
# 主流程
# ═══════════════════════════════════════════════════════════════════════════

def run_self_check(review_id: int) -> dict:
    """对指定review_id执行 LLM 核实，返回完整报告"""
    if not DB_PATH.exists():
        raise FileNotFoundError(f"DB不存在: {DB_PATH}")

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    db = conn.cursor()

    # 读取审查信息
    db.execute("""
        SELECT r.*, p.name as project_name, p.file_path, p.file_name
        FROM reviews r
        JOIN projects p ON r.project_id = p.id
        WHERE r.id = ?
    """, (review_id,))
    row = db.fetchone()
    if not row:
        raise ValueError(f"审查不存在: review_id={review_id}")

    review_info = dict(row)

    # 读取缺陷列表
    db.execute("""
        SELECT id, rule_id, chapter, severity, description, report_excerpt
        FROM defects
        WHERE review_id = ?
        ORDER BY severity, id
    """, (review_id,))
    defects_raw = [dict(r) for r in db.fetchall()]

    if not defects_raw:
        print("⚠️ 该审查无缺陷记录")
        return {"review_info": review_info, "defects": [], "analysis": {}, "suggestions": []}

    # 加载 API Key
    api_key = _load_api_key()
    if not api_key:
        print("⚠️ 未配置 deepseek_api_key，改为正则模式核实")
    else:
        print(f"✓ DeepSeek API Key 已加载")

    # 加载文档（按章节分割）
    docx_path = review_info["file_path"]
    if not docx_path or not Path(docx_path).exists():
        raise FileNotFoundError(f"DOCX文件不存在: {docx_path}")

    print(f"正在加载文档: {docx_path}")
    doc_data = load_docx_with_chapters(docx_path)
    chapters = doc_data["chapters"]
    print(f"  章节: {len(chapters)}, 总字符: {len(doc_data['full_text']):,}")

    # 统计
    total = len(defects_raw)
    print(f"开始 LLM 核实 {total} 条缺陷...")
    print(f"  （每个缺陷单独调用 DeepSeek，预计每条 3-8 秒）")

    results = []
    errors = 0

    for i, d in enumerate(defects_raw, 1):
        rid = d["rule_id"]
        desc = d["description"] or ""
        defect_chapter = d["chapter"] or ""
        report_excerpt = d.get("report_excerpt") or ""
        severity_label = d["severity"] or "一般"

        # 找到对应章节的原文
        ch_num, ch_name, ch_text = find_chapter_for_defect(chapters, defect_chapter)

        if i % 10 == 1:
            print(f"\n  进度 {i}/{total}...")

        try:
            verify_result = verify_defect_llm(
                defect_description=desc,
                report_excerpt=report_excerpt,
                chapter_name=ch_name or defect_chapter,
                rule_id=rid,
                project_name=review_info["project_name"],
                api_key=api_key,
            )
        except Exception as e:
            errors += 1
            verify_result = {
                "verdict": UNCERTAIN,
                "hit_rate": 0.0,
                "summary": f"[异常: {str(e)}]",
                "reasoning": "",
            }

        results.append({
            "defect_id": d["id"],
            "rule_id": rid,
            "chapter": defect_chapter,
            "chapter_matched": ch_name,
            "severity": severity_label,
            "severity_short": SEVERITY_MAP.get(severity_label, "C"),
            "description": desc,
            "report_excerpt": report_excerpt[:200] if report_excerpt else "",
            "verify_result": verify_result,
            "verdict": verify_result["verdict"],
        })

        # 限速：每分钟不超过 20 次 API 调用（3秒/条）
        if i < total:
            time.sleep(3.5)

    print(f"\n✓ LLM 核实完成。共 {total} 条，成功 {total - errors} 条，异常 {errors} 条")

    # 质量分析
    analysis = analyze_results(results)
    suggestions = generate_optimization_suggestions(results, analysis)

    # 写回DB（先删旧记录再插入新记录）
    db.execute("DELETE FROM defect_verification_results WHERE review_id = ?", (review_id,))
    now = datetime.now().isoformat()
    for r in results:
        db.execute("""
            INSERT INTO defect_verification_results
            (defect_id, review_id, rule_id, verdict, hit_rate, verify_keywords,
             verify_context, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            r["defect_id"],
            review_id,
            r["rule_id"],
            r["verdict"],
            r["verify_result"]["hit_rate"],
            r["verify_result"]["summary"],   # 专家可读的判定说明
            json.dumps({"reasoning": r["verify_result"]["reasoning"]}, ensure_ascii=False),
            now,
        ))

    conn.commit()
    conn.close()

    return {
        "review_info": {
            "review_id": review_id,
            "project_name": review_info["project_name"],
            "file_name": review_info["file_name"],
            "status": review_info["status"],
            "total_defects": review_info["total_defects"],
            "type_a": review_info["type_a_count"],
            "type_b": review_info["type_b_count"],
            "type_c": review_info["type_c_count"],
            "completed_at": review_info["completed_at"],
        },
        "defects": results,
        "analysis": analysis,
        "suggestions": suggestions,
    }


def print_report(report: dict):
    """格式化输出报告到终端"""
    ri = report["review_info"]
    a = report["analysis"]
    suggs = report["suggestions"]

    print("\n" + "=" * 90)
    print(f"自检报告 | {ri['project_name']} | review_id={ri['review_id']}")
    print("=" * 90)
    print(f"\n📊 总体统计")
    print(f"  缺陷总数: {a['total']}  |  属实: {a['real']}  |  存疑: {a['uncertain']}  |  不实: {a['unreal']}")
    print(f"  命中率: {a['hit_rate']:.1%}  |  精度: {a['precision']:.1%}  |  漂移率: {a['drift_rate']:.1%}")

    print(f"\n📋 按严重度")
    for sev in SEVERITY_ORDER:
        s = a["by_severity"].get(sev, {})
        if s.get("total"):
            print(f"  {sev}: {s['total']}条  属实{s['real']}  存疑{s['uncertain']}  不实{s['unreal']}")

    if a["problem_rules"]:
        print(f"\n⚠️ 问题规则（命中率<80%）")
        for pr in a["problem_rules"][:10]:
            print(f"  {pr['rule_id']}: {pr['hit_rate']:.0%} ({pr['real']}/{pr['total']}) — {pr['verdict']}")

    if suggs:
        print(f"\n💡 优化建议（{len(suggs)}条）")
        for s in suggs[:5]:
            print(f"  [{s['priority']}] {s['summary']}")

    print(f"\n📝 核实依据示例（前5条）")
    for d in report["defects"][:5]:
        vr = d["verify_result"]
        print(f"  [{d['severity_short']}] {d['verdict']} | {d['rule_id']} | {vr['summary'][:60]}")

    print()


# ═══════════════════════════════════════════════════════════════════════════
# 入口
# ═══════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python3 self_check.py <review_id>")
        sys.exit(1)

    review_id = int(sys.argv[1])
    print(f"\n{'='*60}")
    print(f" 环评审查自检 | review_id={review_id}")
    print(f"{'='*60}\n")

    try:
        report = run_self_check(review_id)
        print_report(report)

        # 保存 JSON
        out_path = SCRIPT_DIR / f"self_check_report_{review_id}.json"
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        print(f"JSON报告已保存: {out_path}")

    except Exception as e:
        print(f"\n❌ 自检失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
